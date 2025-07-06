#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os

def create_admin_guide():
    # Tạo document mới
    doc = Document()
    
    # Thiết lập style cho tiêu đề
    styles = doc.styles
    title_style = styles['Title']
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    heading1_style = styles['Heading 1']
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True
    
    heading2_style = styles['Heading 2']
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True
    
    # Tiêu đề chính
    title = doc.add_heading('HƯỚNG DẪN SỬ DỤNG TRANG ADMIN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thông tin chung
    doc.add_heading('1. TỔNG QUAN', level=1)
    doc.add_paragraph('Trang Admin là nơi quản lý toàn bộ nội dung và cài đặt của website. Từ đây, bạn có thể:')
    
    features = doc.add_paragraph()
    features.add_run('• Quản lý tin tức (thêm, sửa, xóa)\n').bold = False
    features.add_run('• Quản lý banner quảng cáo\n').bold = False
    features.add_run('• Quản lý thư viện ảnh\n').bold = False
    features.add_run('• Quản lý đối tác\n').bold = False
    features.add_run('• Xem và phản hồi liên hệ từ khách hàng\n').bold = False
    features.add_run('• Quản lý người dùng\n').bold = False
    features.add_run('• Cài đặt hệ thống\n').bold = False
    
    # Thêm placeholder cho ảnh dashboard
    doc.add_paragraph('Hình ảnh: Dashboard tổng quan')
    doc.add_paragraph('[CHÈN ẢNH DASHBOARD Ở ĐÂY]')
    
    # Đăng nhập
    doc.add_heading('2. ĐĂNG NHẬP VÀO TRANG ADMIN', level=1)
    doc.add_paragraph('Để truy cập trang admin, bạn cần:')
    
    login_steps = doc.add_paragraph()
    login_steps.add_run('1. Truy cập đường dẫn: ').bold = True
    login_steps.add_run('/admin/login\n').bold = False
    login_steps.add_run('2. Nhập tên đăng nhập và mật khẩu\n').bold = False
    login_steps.add_run('3. Nhấn nút "Đăng nhập"\n').bold = False
    
    doc.add_paragraph('Hình ảnh: Trang đăng nhập')
    doc.add_paragraph('[CHÈN ẢNH TRANG ĐĂNG NHẬP Ở ĐÂY]')
    
    # Quản lý tin tức
    doc.add_heading('3. QUẢN LÝ TIN TỨC', level=1)
    
    doc.add_heading('3.1. Xem danh sách tin tức', level=2)
    doc.add_paragraph('Từ menu chính, chọn "Tin tức" để xem danh sách tất cả tin tức.')
    doc.add_paragraph('Hình ảnh: Danh sách tin tức')
    doc.add_paragraph('[CHÈN ẢNH DANH SÁCH TIN TỨC Ở ĐÂY]')
    
    doc.add_heading('3.2. Thêm tin tức mới', level=2)
    add_news_steps = doc.add_paragraph()
    add_news_steps.add_run('1. Từ trang danh sách tin tức, nhấn nút "Thêm mới"\n').bold = False
    add_news_steps.add_run('2. Điền thông tin tin tức:\n').bold = False
    add_news_steps.add_run('   • Tiêu đề: Nhập tiêu đề tin tức\n').bold = False
    add_news_steps.add_run('   • Mô tả ngắn: Tóm tắt nội dung\n').bold = False
    add_news_steps.add_run('   • Nội dung: Sử dụng trình soạn thảo để nhập nội dung chi tiết\n').bold = False
    add_news_steps.add_run('   • Ảnh đại diện: Tải lên ảnh minh họa\n').bold = False
    add_news_steps.add_run('   • Trạng thái: Chọn "Hiển thị" hoặc "Ẩn"\n').bold = False
    add_news_steps.add_run('3. Nhấn "Lưu" để lưu tin tức\n').bold = False
    
    doc.add_paragraph('Hình ảnh: Form thêm tin tức')
    doc.add_paragraph('[CHÈN ẢNH FORM THÊM TIN TỨC Ở ĐÂY]')
    
    doc.add_heading('3.3. Chỉnh sửa tin tức', level=2)
    edit_news_steps = doc.add_paragraph()
    edit_news_steps.add_run('1. Từ danh sách tin tức, nhấn nút "Sửa" bên cạnh tin tức cần chỉnh sửa\n').bold = False
    edit_news_steps.add_run('2. Thay đổi thông tin cần thiết\n').bold = False
    edit_news_steps.add_run('3. Nhấn "Cập nhật" để lưu thay đổi\n').bold = False
    
    doc.add_paragraph('Hình ảnh: Form chỉnh sửa tin tức')
    doc.add_paragraph('[CHÈN ẢNH FORM CHỈNH SỬA TIN TỨC Ở ĐÂY]')
    
    doc.add_heading('3.4. Xóa tin tức', level=2)
    delete_news_steps = doc.add_paragraph()
    delete_news_steps.add_run('1. Từ danh sách tin tức, nhấn nút "Xóa" bên cạnh tin tức cần xóa\n').bold = False
    delete_news_steps.add_run('2. Xác nhận xóa trong hộp thoại hiện ra\n').bold = False
    delete_news_steps.add_run('3. Tin tức sẽ bị xóa vĩnh viễn\n').bold = False
    
    # Quản lý banner
    doc.add_heading('4. QUẢN LÝ BANNER', level=1)
    
    doc.add_heading('4.1. Xem danh sách banner', level=2)
    doc.add_paragraph('Từ menu chính, chọn "Banner" để xem danh sách tất cả banner.')
    doc.add_paragraph('Hình ảnh: Danh sách banner')
    doc.add_paragraph('[CHÈN ẢNH DANH SÁCH BANNER Ở ĐÂY]')
    
    doc.add_heading('4.2. Thêm banner mới', level=2)
    add_banner_steps = doc.add_paragraph()
    add_banner_steps.add_run('1. Từ trang danh sách banner, nhấn nút "Thêm mới"\n').bold = False
    add_banner_steps.add_run('2. Điền thông tin banner:\n').bold = False
    add_banner_steps.add_run('   • Tiêu đề: Nhập tiêu đề banner\n').bold = False
    add_banner_steps.add_run('   • Mô tả: Mô tả ngắn về banner\n').bold = False
    add_banner_steps.add_run('   • Ảnh: Tải lên ảnh banner\n').bold = False
    add_banner_steps.add_run('   • Liên kết: URL khi click vào banner (tùy chọn)\n').bold = False
    add_banner_steps.add_run('   • Thứ tự: Số thứ tự hiển thị\n').bold = False
    add_banner_steps.add_run('   • Trạng thái: Chọn "Hiển thị" hoặc "Ẩn"\n').bold = False
    add_banner_steps.add_run('3. Nhấn "Lưu" để lưu banner\n').bold = False
    
    doc.add_paragraph('Hình ảnh: Form thêm banner')
    doc.add_paragraph('[CHÈN ẢNH FORM THÊM BANNER Ở ĐÂY]')
    
    doc.add_heading('4.3. Chỉnh sửa banner', level=2)
    doc.add_paragraph('Tương tự như thêm mới, nhưng thông tin sẽ được điền sẵn để chỉnh sửa.')
    doc.add_paragraph('Hình ảnh: Form chỉnh sửa banner')
    doc.add_paragraph('[CHÈN ẢNH FORM CHỈNH SỬA BANNER Ở ĐÂY]')
    
    doc.add_heading('4.4. Xóa banner', level=2)
    doc.add_paragraph('Nhấn nút "Xóa" và xác nhận để xóa banner.')
    
    # Quản lý thư viện ảnh
    doc.add_heading('5. QUẢN LÝ THƯ VIỆN ẢNH', level=1)
    
    doc.add_heading('5.1. Xem danh sách ảnh', level=2)
    doc.add_paragraph('Từ menu chính, chọn "Thư viện ảnh" để xem tất cả ảnh.')
    doc.add_paragraph('Hình ảnh: Thư viện ảnh')
    doc.add_paragraph('[CHÈN ẢNH THƯ VIỆN ẢNH Ở ĐÂY]')
    
    doc.add_heading('5.2. Thêm ảnh mới', level=2)
    add_gallery_steps = doc.add_paragraph()
    add_gallery_steps.add_run('1. Từ trang thư viện ảnh, nhấn nút "Thêm mới"\n').bold = False
    add_gallery_steps.add_run('2. Điền thông tin ảnh:\n').bold = False
    add_gallery_steps.add_run('   • Tiêu đề: Nhập tiêu đề ảnh\n').bold = False
    add_gallery_steps.add_run('   • Mô tả: Mô tả về ảnh\n').bold = False
    add_gallery_steps.add_run('   • Ảnh: Tải lên file ảnh\n').bold = False
    add_gallery_steps.add_run('   • Danh mục: Chọn danh mục ảnh (nếu có)\n').bold = False
    add_gallery_steps.add_run('   • Trạng thái: Chọn "Hiển thị" hoặc "Ẩn"\n').bold = False
    add_gallery_steps.add_run('3. Nhấn "Lưu" để lưu ảnh\n').bold = False
    
    doc.add_paragraph('Hình ảnh: Form thêm ảnh')
    doc.add_paragraph('[CHÈN ẢNH FORM THÊM ẢNH Ở ĐÂY]')
    
    doc.add_heading('5.3. Chỉnh sửa ảnh', level=2)
    doc.add_paragraph('Nhấn nút "Sửa" để chỉnh sửa thông tin ảnh.')
    doc.add_paragraph('Hình ảnh: Form chỉnh sửa ảnh')
    doc.add_paragraph('[CHÈN ẢNH FORM CHỈNH SỬA ẢNH Ở ĐÂY]')
    
    doc.add_heading('5.4. Xóa ảnh', level=2)
    doc.add_paragraph('Nhấn nút "Xóa" và xác nhận để xóa ảnh.')
    
    # Quản lý đối tác
    doc.add_heading('6. QUẢN LÝ ĐỐI TÁC', level=1)
    
    doc.add_heading('6.1. Xem danh sách đối tác', level=2)
    doc.add_paragraph('Từ menu chính, chọn "Đối tác" để xem danh sách đối tác.')
    doc.add_paragraph('Hình ảnh: Danh sách đối tác')
    doc.add_paragraph('[CHÈN ẢNH DANH SÁCH ĐỐI TÁC Ở ĐÂY]')
    
    doc.add_heading('6.2. Thêm đối tác mới', level=2)
    add_partner_steps = doc.add_paragraph()
    add_partner_steps.add_run('1. Từ trang danh sách đối tác, nhấn nút "Thêm mới"\n').bold = False
    add_partner_steps.add_run('2. Điền thông tin đối tác:\n').bold = False
    add_partner_steps.add_run('   • Tên đối tác: Nhập tên công ty/đối tác\n').bold = False
    add_partner_steps.add_run('   • Logo: Tải lên logo đối tác\n').bold = False
    add_partner_steps.add_run('   • Website: URL website đối tác (tùy chọn)\n').bold = False
    add_partner_steps.add_run('   • Mô tả: Thông tin về đối tác\n').bold = False
    add_partner_steps.add_run('   • Thứ tự: Số thứ tự hiển thị\n').bold = False
    add_partner_steps.add_run('   • Trạng thái: Chọn "Hiển thị" hoặc "Ẩn"\n').bold = False
    add_partner_steps.add_run('3. Nhấn "Lưu" để lưu đối tác\n').bold = False
    
    doc.add_paragraph('Hình ảnh: Form thêm đối tác')
    doc.add_paragraph('[CHÈN ẢNH FORM THÊM ĐỐI TÁC Ở ĐÂY]')
    
    doc.add_heading('6.3. Chỉnh sửa đối tác', level=2)
    doc.add_paragraph('Nhấn nút "Sửa" để chỉnh sửa thông tin đối tác.')
    doc.add_paragraph('Hình ảnh: Form chỉnh sửa đối tác')
    doc.add_paragraph('[CHÈN ẢNH FORM CHỈNH SỬA ĐỐI TÁC Ở ĐÂY]')
    
    doc.add_heading('6.4. Xóa đối tác', level=2)
    doc.add_paragraph('Nhấn nút "Xóa" và xác nhận để xóa đối tác.')
    
    # Quản lý liên hệ
    doc.add_heading('7. QUẢN LÝ LIÊN HỆ', level=1)
    
    doc.add_heading('7.1. Xem danh sách liên hệ', level=2)
    doc.add_paragraph('Từ menu chính, chọn "Liên hệ" để xem tất cả tin nhắn từ khách hàng.')
    doc.add_paragraph('Hình ảnh: Danh sách liên hệ')
    doc.add_paragraph('[CHÈN ẢNH DANH SÁCH LIÊN HỆ Ở ĐÂY]')
    
    doc.add_heading('7.2. Xem chi tiết liên hệ', level=2)
    doc.add_paragraph('Nhấn vào tiêu đề liên hệ để xem nội dung chi tiết.')
    doc.add_paragraph('Hình ảnh: Chi tiết liên hệ')
    doc.add_paragraph('[CHÈN ẢNH CHI TIẾT LIÊN HỆ Ở ĐÂY]')
    
    doc.add_heading('7.3. Xóa liên hệ', level=2)
    doc.add_paragraph('Nhấn nút "Xóa" và xác nhận để xóa liên hệ.')
    
    # Quản lý người dùng
    doc.add_heading('8. QUẢN LÝ NGƯỜI DÙNG', level=1)
    
    doc.add_heading('8.1. Xem danh sách người dùng', level=2)
    doc.add_paragraph('Từ menu chính, chọn "Người dùng" để xem danh sách tất cả người dùng.')
    doc.add_paragraph('Hình ảnh: Danh sách người dùng')
    doc.add_paragraph('[CHÈN ẢNH DANH SÁCH NGƯỜI DÙNG Ở ĐÂY]')
    
    doc.add_heading('8.2. Thêm người dùng mới', level=2)
    add_user_steps = doc.add_paragraph()
    add_user_steps.add_run('1. Từ trang danh sách người dùng, nhấn nút "Thêm mới"\n').bold = False
    add_user_steps.add_run('2. Điền thông tin người dùng:\n').bold = False
    add_user_steps.add_run('   • Tên: Họ và tên người dùng\n').bold = False
    add_user_steps.add_run('   • Email: Địa chỉ email\n').bold = False
    add_user_steps.add_run('   • Mật khẩu: Mật khẩu đăng nhập\n').bold = False
    add_user_steps.add_run('   • Vai trò: Chọn quyền hạn\n').bold = False
    add_user_steps.add_run('3. Nhấn "Lưu" để tạo người dùng\n').bold = False
    
    doc.add_paragraph('Hình ảnh: Form thêm người dùng')
    doc.add_paragraph('[CHÈN ẢNH FORM THÊM NGƯỜI DÙNG Ở ĐÂY]')
    
    doc.add_heading('8.3. Chỉnh sửa người dùng', level=2)
    doc.add_paragraph('Nhấn nút "Sửa" để chỉnh sửa thông tin người dùng.')
    doc.add_paragraph('Hình ảnh: Form chỉnh sửa người dùng')
    doc.add_paragraph('[CHÈN ẢNH FORM CHỈNH SỬA NGƯỜI DÙNG Ở ĐÂY]')
    
    doc.add_heading('8.4. Xóa người dùng', level=2)
    doc.add_paragraph('Nhấn nút "Xóa" và xác nhận để xóa người dùng.')
    
    # Cài đặt hệ thống
    doc.add_heading('9. CÀI ĐẶT HỆ THỐNG', level=1)
    
    doc.add_heading('9.1. Cài đặt chung', level=2)
    doc.add_paragraph('Từ menu chính, chọn "Cài đặt" để thay đổi cấu hình hệ thống.')
    doc.add_paragraph('Hình ảnh: Trang cài đặt')
    doc.add_paragraph('[CHÈN ẢNH TRANG CÀI ĐẶT Ở ĐÂY]')
    
    doc.add_heading('9.2. Cài đặt thông tin website', level=2)
    settings_steps = doc.add_paragraph()
    settings_steps.add_run('1. Điền thông tin website:\n').bold = False
    settings_steps.add_run('   • Tên website\n').bold = False
    settings_steps.add_run('   • Mô tả website\n').bold = False
    settings_steps.add_run('   • Logo website\n').bold = False
    settings_steps.add_run('   • Favicon\n').bold = False
    settings_steps.add_run('2. Nhấn "Lưu" để cập nhật\n').bold = False
    
    doc.add_paragraph('Hình ảnh: Cài đặt thông tin website')
    doc.add_paragraph('[CHÈN ẢNH CÀI ĐẶT THÔNG TIN WEBSITE Ở ĐÂY]')
    
    doc.add_heading('9.3. Cài đặt liên hệ', level=2)
    contact_settings_steps = doc.add_paragraph()
    contact_settings_steps.add_run('1. Điền thông tin liên hệ:\n').bold = False
    contact_settings_steps.add_run('   • Địa chỉ\n').bold = False
    contact_settings_steps.add_run('   • Số điện thoại\n').bold = False
    contact_settings_steps.add_run('   • Email\n').bold = False
    contact_settings_steps.add_run('   • Giờ làm việc\n').bold = False
    contact_settings_steps.add_run('2. Nhấn "Lưu" để cập nhật\n').bold = False
    
    doc.add_paragraph('Hình ảnh: Cài đặt thông tin liên hệ')
    doc.add_paragraph('[CHÈN ẢNH CÀI ĐẶT THÔNG TIN LIÊN HỆ Ở ĐÂY]')
    
    # Lưu ý quan trọng
    doc.add_heading('10. LƯU Ý QUAN TRỌNG', level=1)
    
    important_notes = doc.add_paragraph()
    important_notes.add_run('• Luôn sao lưu dữ liệu trước khi thực hiện các thao tác quan trọng\n').bold = False
    important_notes.add_run('• Kiểm tra kỹ nội dung trước khi đăng tin tức\n').bold = False
    important_notes.add_run('• Sử dụng ảnh có kích thước phù hợp để tối ưu tốc độ tải trang\n').bold = False
    important_notes.add_run('• Thay đổi mật khẩu định kỳ để bảo mật\n').bold = False
    important_notes.add_run('• Không chia sẻ thông tin đăng nhập với người khác\n').bold = False
    important_notes.add_run('• Liên hệ quản trị viên nếu gặp vấn đề kỹ thuật\n').bold = False
    
    # Hỗ trợ
    doc.add_heading('11. HỖ TRỢ', level=1)
    doc.add_paragraph('Nếu bạn gặp khó khăn trong quá trình sử dụng, vui lòng liên hệ:')
    doc.add_paragraph('• Email: support@example.com')
    doc.add_paragraph('• Điện thoại: 0123 456 789')
    doc.add_paragraph('• Thời gian hỗ trợ: 8:00 - 17:00 (Thứ 2 - Thứ 6)')
    
    # Quản lý trang
    doc.add_heading('4. QUẢN LÝ TRANG', level=1)
    
    doc.add_heading('4.1. Xem danh sách trang', level=2)
    doc.add_paragraph('Từ menu chính, chọn "Quản lý trang" để xem danh sách tất cả các trang nội dung như Trang chủ, Trang giới thiệu, Thông tin chung, v.v.')
    doc.add_paragraph('Hình ảnh: Danh sách các trang')
    doc.add_paragraph('[CHÈN ẢNH DANH SÁCH TRANG Ở ĐÂY]')
    
    doc.add_heading('4.2. Thêm trang mới', level=2)
    add_page_steps = doc.add_paragraph()
    add_page_steps.add_run('1. Từ trang danh sách trang, nhấn nút "Thêm mới"\n').bold = False
    add_page_steps.add_run('2. Điền thông tin trang:\n').bold = False
    add_page_steps.add_run('   • Tiêu đề: Nhập tiêu đề trang\n').bold = False
    add_page_steps.add_run('   • Nội dung: Sử dụng trình soạn thảo để nhập nội dung chi tiết\n').bold = False
    add_page_steps.add_run('   • Trạng thái: Chọn "Hiển thị" hoặc "Ẩn"\n').bold = False
    add_page_steps.add_run('3. Nhấn "Lưu" để tạo trang mới\n').bold = False
    doc.add_paragraph('Hình ảnh: Form thêm trang')
    doc.add_paragraph('[CHÈN ẢNH FORM THÊM TRANG Ở ĐÂY]')
    
    doc.add_heading('4.3. Chỉnh sửa trang', level=2)
    edit_page_steps = doc.add_paragraph()
    edit_page_steps.add_run('1. Từ danh sách trang, nhấn nút "Sửa" bên cạnh trang cần chỉnh sửa\n').bold = False
    edit_page_steps.add_run('2. Thay đổi thông tin cần thiết\n').bold = False
    edit_page_steps.add_run('3. Nhấn "Cập nhật" để lưu thay đổi\n').bold = False
    doc.add_paragraph('Hình ảnh: Form chỉnh sửa trang')
    doc.add_paragraph('[CHÈN ẢNH FORM CHỈNH SỬA TRANG Ở ĐÂY]')
    
    doc.add_heading('4.4. Xóa trang', level=2)
    delete_page_steps = doc.add_paragraph()
    delete_page_steps.add_run('1. Từ danh sách trang, nhấn nút "Xóa" bên cạnh trang cần xóa\n').bold = False
    delete_page_steps.add_run('2. Xác nhận xóa trong hộp thoại hiện ra\n').bold = False
    delete_page_steps.add_run('3. Trang sẽ bị xóa vĩnh viễn\n').bold = False
    
    # Lưu file
    doc.save('Huong_dan_su_dung_Admin.docx')
    print("Đã tạo file hướng dẫn: Huong_dan_su_dung_Admin.docx")

if __name__ == "__main__":
    create_admin_guide() 